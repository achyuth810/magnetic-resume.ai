# utils/docx_writer.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.templates import get_template


def write_resume_docx(text: str, output_path: str, title: str = "TAILORED RESUME", template: str = "ATS_CLASSIC"):
    config = get_template(template)
    styles = config["styles"]

    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(config["left_margin"] / 72)
    section.right_margin = Inches(config["right_margin"] / 72)
    section.top_margin = Inches(config["top_margin"] / 72)
    section.bottom_margin = Inches(config["bottom_margin"] / 72)

    # Name
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    
    if lines:
        name_para = doc.add_paragraph(lines[0])
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.name = config["bold_font"]
        name_run.font.size = Pt(config["name_font_size"])
        name_run.bold = True

    # Contact line
    if len(lines) > 1:
        contact_para = doc.add_paragraph(lines[1])
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.runs[0]
        contact_run.font.name = config["font_name"]
        contact_run.font.size = Pt(config["small_font_size"] + 0.5)

    doc.add_paragraph()  # spacing

    # Rest of the content
    i = 2
    while i < len(lines):
        line = lines[i]
        
        if line.isupper() and len(line) > 3:  # Section header
            para = doc.add_paragraph(line)
            para.style = "Heading 1"
            run = para.runs[0]
            run.font.name = config["bold_font"]
            run.font.size = Pt(config["header_font_size"])
            run.font.color.rgb = config["section_color"].rgb if hasattr(config["section_color"], "rgb") else None
            doc.add_paragraph()  # spacing
            
        elif line.startswith("- ") or line.startswith("• "):
            para = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            run = para.runs[0]
            run.font.name = config["font_name"]
            run.font.size = Pt(config["body_font_size"])
            
        else:
            para = doc.add_paragraph(line)
            run = para.runs[0]
            run.font.name = config["font_name"]
            run.font.size = Pt(config["body_font_size"])
        
        i += 1

    doc.save(output_path)